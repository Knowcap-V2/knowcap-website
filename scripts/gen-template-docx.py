#!/usr/bin/env python3
"""
gen-template-docx.py — turn a plain-text copy-paste template into a real
downloadable .docx, for the blogger routine's "free template" posts.

Reads a text file (the same content that used to go into a broken
```-fenced code block on the blog page) and structures it into a Word
document: ALL-CAPS lines become section headings, "- " lines become
bullets, everything else is a body paragraph (blank underscores render
fine in any font, no monospace needed).

Usage:
  python3 scripts/gen-template-docx.py \
    --title "Meeting Minutes Template" \
    --subtitle "Free copy-paste template - knowcap.ai/blog/how-to-do-meeting-minutes" \
    --text-file /tmp/template.txt \
    --out app/public/downloads/meeting-minutes-template.docx

Requires: pip install python-docx
"""
import argparse
import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ACCENT = RGBColor(0x1F, 0x6B, 0x3A)
INK = RGBColor(0x18, 0x18, 0x1B)
INK_2 = RGBColor(0x4A, 0x4F, 0x5A)


def apply_rtl(paragraph):
    """Right-align + set bidi paragraph direction (python-docx has no high-level API for this)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = rPr.makeelement(qn('w:rtl'), {})
        rPr.append(rtl)


def is_section_heading(line: str) -> bool:
    s = line.strip()
    if len(s) <= 1:
        return False
    # Bare label ending in a colon (e.g. Arabic "جدول الأعمال:") — no cased
    # script needed, works for any language.
    if s.endswith(':') or s.endswith('：'):
        return True
    # ALL-CAPS heading in a cased script (e.g. "MEETING MINUTES"). str.isupper()
    # is False for scripts with no case distinction (Arabic, CJK, digits-only),
    # so this never false-positives on non-Latin text the way `s == s.upper()`
    # would (that comparison is a no-op — and therefore always true — on
    # uncased text).
    return s.isupper()


def build(title: str, subtitle: str, text: str, out_path: str, source_url: str, rtl: bool = False):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_paragraph()
    run = h.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = INK
    if rtl:
        apply_rtl(h)

    if subtitle:
        sub = doc.add_paragraph()
        srun = sub.add_run(subtitle)
        srun.italic = True
        srun.font.size = Pt(10.5)
        srun.font.color.rgb = INK_2
        if rtl:
            apply_rtl(sub)

    doc.add_paragraph()

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        if is_section_heading(line):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = ACCENT
            p.space_before = Pt(10)
            if rtl:
                apply_rtl(p)
            continue
        if line.strip().startswith("- "):
            p = doc.add_paragraph(line.strip()[2:], style="List Bullet")
            if rtl:
                apply_rtl(p)
            continue
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.color.rgb = INK_2
        if rtl:
            apply_rtl(p)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = foot.add_run(f"Free template from Knowcap — {source_url}")
    frun.font.size = Pt(9)
    frun.font.color.rgb = INK_2

    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    print(f"OK  {out_path}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--title", required=True)
    ap.add_argument("--subtitle", default="")
    ap.add_argument("--text-file", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--source-url", default="knowcap.ai/blog")
    ap.add_argument("--rtl", action="store_true", help="right-align + bidi paragraph direction (Arabic/Hebrew)")
    args = ap.parse_args()

    with open(args.text_file, "r", encoding="utf-8") as f:
        text = f.read()

    build(args.title, args.subtitle, text, args.out, args.source_url, rtl=args.rtl)


if __name__ == "__main__":
    main()
